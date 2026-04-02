"""
Report generation module for Graduate Outcomes reports.

All three raw source tables plus master DB are consulted for every section.
Qualtrics, LinkedIn, and Clearinghouse payloads are all read directly.

  Career outcomes   – Qualtrics STATUS → LinkedIn → Clearinghouse → master DB → Unresolved
  Geographic        – Qualtrics EMP_CITY1_1 + LinkedIn employer_state
  Nature/modality   – Qualtrics EMP_NATURE/EMP_FIELD/EMP_JOBSITE + EMP_TYPE (emp status) + LinkedIn employment_modality
  Salary            – Qualtrics EMP_SALARY/EMP_BONUS (no salary field in LinkedIn/CH)
  Search methods    – Qualtrics EMP_HOW_* (survey-specific, not in other sources)
  Business          – Qualtrics STBUS_* + LinkedIn started_business_*
  Volunteer         – Qualtrics VOL_* + LinkedIn volunteer_*
  Military          – Qualtrics MIL_* + LinkedIn joined_military_branch
  Continuing ed     – Qualtrics CONTEDU_* + Clearinghouse + LinkedIn CE fields
  Internships       – Qualtrics NUMINTERN/*_INT_* (survey-specific)
  OOC experience    – Qualtrics OTHEREXP_* (survey-specific)
  Appendix A        – Qualtrics EMP_ORG_1/TITLE + LinkedIn employer/job_title
  Appendix B        – Qualtrics CONTEDU + Clearinghouse + LinkedIn CE programs
"""

import io
import re
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import database

# ── Lookup tables ──────────────────────────────────────────────────────────────

STATUS_TO_OUTCOME = {
    "employed full-time":                              "Employed full-time",
    "employed part-time":                              "Employed part-time",
    "accepted into a program of continuing education": "Continuing education",
    # "applied to graduate school" = intent without acceptance → Unplaced per spec
    "applied to graduate school":                      "Unplaced",
    "starting my own business":                        "Starting a business",
    "serving in the u.s. armed forces":                "Serving in the U.S. Armed Forces",
    "participating in a service":                      "Volunteering or service program",
    "not seeking":                                     "NOT seeking",
    "actively seeking":                                "Unplaced",
}

EMP_HOW_LABEL = {
    "NONUMDSITE": "Non-UMD job site / company site / social media",
    "OTHER":      "Other",
    "FAMILY":     "Family/Friends contacts",
    "INTERN":     "Internship",
    "CURRENT":    "Currently employed with organization",
    "UMDSITE":    "UMD online job site (e.g. Handshake, Careers4Terps)",
    "FACULTY":    "Faculty/Staff contacts",
    "UMDCAREER":  "UMD Career Fairs",
    "ALUM":       "Alumni network",
    "FRIENDS":    "Friends/Colleagues",
}

OTHEREXP_LABEL = {
    "OTHEREXP_77": "Part-time employment – on or off campus",
    "OTHEREXP_78": "Full-time employment – on or off campus",
    "OTHEREXP_81": "Research program(s) or projects",
    "OTHEREXP_83": "Early field experience / student teaching",
    "OTHEREXP_84": "Clinical or hospital experience",
    "OTHEREXP_86": "Study or work abroad",
    "OTHEREXP_89": "Community service-learning / volunteer",
    "OTHEREXP_90": "Student group leadership",
    "OTHEREXP_91": "Student group membership",
    "OTHEREXP_92": "Terrapin Teachers",
    "OTHEREXP_94": "Other experience",
    "OTHEREXP_95": "None of the above",
}

STATE_ABBR = {
    "AL": "Alabama",    "AK": "Alaska",    "AZ": "Arizona",    "AR": "Arkansas",
    "CA": "California", "CO": "Colorado",  "CT": "Connecticut","DE": "Delaware",
    "DC": "District of Columbia", "FL": "Florida", "GA": "Georgia", "HI": "Hawaii",
    "ID": "Idaho",   "IL": "Illinois", "IN": "Indiana",  "IA": "Iowa",
    "KS": "Kansas",  "KY": "Kentucky", "LA": "Louisiana", "ME": "Maine",
    "MD": "Maryland","MA": "Massachusetts","MI": "Michigan","MN": "Minnesota",
    "MS": "Mississippi","MO": "Missouri","MT": "Montana","NE": "Nebraska",
    "NV": "Nevada",  "NH": "New Hampshire","NJ": "New Jersey","NM": "New Mexico",
    "NY": "New York","NC": "North Carolina","ND": "North Dakota","OH": "Ohio",
    "OK": "Oklahoma","OR": "Oregon","PA": "Pennsylvania","RI": "Rhode Island",
    "SC": "South Carolina","SD": "South Dakota","TN": "Tennessee","TX": "Texas",
    "UT": "Utah","VT": "Vermont","VA": "Virginia","WA": "Washington",
    "WV": "West Virginia","WI": "Wisconsin","WY": "Wyoming",
}

# ── Helpers ───────────────────────────────────────────────────────────────────

def _map_status(status_str: str) -> str:
    s = status_str.lower()
    for key, label in STATUS_TO_OUTCOME.items():
        if key in s:
            return label
    return "Unresolved"


def _parse_salary_midpoint(salary_str: str):
    nums = re.findall(r"[\d,]+", salary_str)
    if len(nums) >= 2:
        lo = int(nums[0].replace(",", ""))
        hi = int(nums[1].replace(",", ""))
        return (lo + hi) / 2
    elif len(nums) == 1:
        return int(nums[0].replace(",", ""))
    return None


def _extract_state_from_city_str(city_str: str) -> str:
    """Extract full state name from 'City, ST, Country' Qualtrics format."""
    parts = [p.strip() for p in city_str.split(",")]
    if len(parts) >= 2:
        candidate = parts[-2].strip()
        if candidate in STATE_ABBR:
            return STATE_ABBR[candidate]
        if candidate == "DC":
            return "District of Columbia"
        for abbr, full in STATE_ABBR.items():
            if candidate.upper() == abbr or candidate.title() == full:
                return full
        last = parts[-1].strip()
        if last.upper() not in ("USA", "U.S.A", "UNITED STATES"):
            return "Outside US"
    return "Unreported"


def _resolve_state(state_abbr_or_full: str, country: str = "") -> str:
    """Resolve a state abbreviation/name to its full name, handling international."""
    s = state_abbr_or_full.strip()
    c = country.strip().upper()
    if not s:
        return "Unreported"
    if c and c not in ("USA", "U.S.A", "US", "UNITED STATES", ""):
        return "Outside US"
    return STATE_ABBR.get(s.upper(), s)


def _pct(n: int, total: int) -> str:
    if total == 0:
        return "N/A"
    return f"{n / total * 100:.1f}%"


def _percentile(sorted_vals, p):
    if not sorted_vals:
        return None
    idx = (len(sorted_vals) - 1) * p / 100
    lo, hi = int(idx), min(int(idx) + 1, len(sorted_vals) - 1)
    return sorted_vals[lo] + (sorted_vals[hi] - sorted_vals[lo]) * (idx - lo)


def _get_term_cutoff(term: str):
    """
    Return the Phase 1 cutoff date for a graduation term.
    Summer  (Aug grads)  → December 1 same year
    Fall    (Dec grads)  → April 1 next year
    Spring  (May grads)  → September 1 same year
    """
    m = re.search(r'(spring|summer|fall|winter)\s+(\d{4})', term.lower())
    if not m:
        return None
    season, year = m.group(1), int(m.group(2))
    if season == "summer":  return datetime(year, 12, 1)
    if season == "fall":    return datetime(year + 1, 4, 1)
    if season == "spring":  return datetime(year, 9, 1)
    if season == "winter":  return datetime(year, 4, 1)
    return None


# ── Per-student multi-source helpers ─────────────────────────────────────────

def _student_outcome(student: dict) -> str:
    """
    Derive the best available outcome for a student, checking all sources in order:
    Qualtrics STATUS → LinkedIn status/employment → Clearinghouse (always CE)
    → master DB outcome_status → date-based Unresolved/Unplaced → Unresolved

    Unresolved vs Unplaced (per spec):
      Unresolved  = responded to survey BEFORE the Phase 1 cutoff date AND no other info
      Unplaced    = responded on/after cutoff AND no outcome, OR reported actively seeking
                    OR reported "applied to graduate school" (intent without acceptance)
    """
    has_survey_response = bool(student.get("qualtrics_data"))

    # 1. Qualtrics — highest fidelity
    if has_survey_response:
        p = student["qualtrics_data"][0]["payload"]
        status = p.get("STATUS", "").strip()
        if status:
            mapped = _map_status(status)
            # Actively seeking / applied-to-grad / known outcome → return directly
            if mapped != "Unresolved":
                return mapped
            # STATUS present but unrecognized → fall through to check other sources

    # 2. LinkedIn
    if student.get("linkedin_data"):
        li = student["linkedin_data"][0]["payload"]
        li_status = (li.get("status") or "").strip()
        if li_status:
            mapped = _map_status(li_status)
            if mapped != "Unresolved":
                return mapped
        # Infer from field presence
        if li.get("name_of_employer") or li.get("job_title"):
            return "Employed full-time"
        if li.get("continuing_education_institution"):
            return "Continuing education"
        if li.get("joined_military_branch"):
            return "Serving in the U.S. Armed Forces"
        if li.get("volunteer_organization"):
            return "Volunteering or service program"
        if li.get("name_of_started_business"):
            return "Starting a business"

    # 3. Clearinghouse — enrollment always means continuing education
    if student.get("clearinghouse_data"):
        return "Continuing education"

    # 4. Master DB (curated record already saved)
    m = student.get("masterData")
    if m and m.get("currentActivity"):
        return m["currentActivity"]

    # 5. No known outcome — apply time-based Unresolved vs Unplaced for survey respondents
    if has_survey_response:
        recorded_at = student["qualtrics_data"][0].get("recorded_at")
        if recorded_at:
            cutoff = _get_term_cutoff(student.get("term", ""))
            if cutoff:
                # Normalize timezone-aware datetime for comparison
                resp = recorded_at.replace(tzinfo=None) if getattr(recorded_at, "tzinfo", None) else recorded_at
                return "Unplaced" if resp >= cutoff else "Unresolved"
        # Survey response present but no date → treated as Phase 1 (Unresolved)
        return "Unresolved"

    return "Unresolved"


def _student_employer_state(student: dict):
    """Return (state_full_name, is_employed) using EMP_STATE (primary) then EMP_CITY1_1 fallback."""
    if student.get("qualtrics_data"):
        p = student["qualtrics_data"][0]["payload"]
        status = p.get("STATUS", "").strip()
        if "employed" in status.lower() and "seeking" not in status.lower():
            emp_state = p.get("EMP_STATE", "").strip()
            if emp_state:
                return _resolve_state(emp_state), True
            city_str = p.get("EMP_CITY1_1", "").strip()
            if city_str:
                return _extract_state_from_city_str(city_str), True

    if student.get("linkedin_data"):
        li = student["linkedin_data"][0]["payload"]
        outcome = _student_outcome(student)
        if "employed" in outcome.lower():
            state   = (li.get("employer_state")   or "").strip()
            country = (li.get("employer_country") or "").strip()
            if state:
                return _resolve_state(state, country), True

    return None, False


# ── Data aggregation ──────────────────────────────────────────────────────────

def aggregate_report_data(major_filter=None, school_filter=None, term_filter=None) -> dict:
    """
    Aggregate all statistics needed for the report.
    Reads raw Qualtrics, LinkedIn, and Clearinghouse payloads plus master DB.
    """

    students = database.get_students_with_data(
        limit=None, offset=None,
        major_filter=major_filter,
        school_filter=school_filter,
        term_filter=term_filter,
    )

    total_graduates = len(students)

    # ── Response / knowledge rates ────────────────────────────────────────────
    with_survey       = [s for s in students if s.get("qualtrics_data")]
    with_linkedin     = [s for s in students if s.get("linkedin_data")]
    with_clearinghouse= [s for s in students if s.get("clearinghouse_data")]
    with_any          = [s for s in students if (
        s.get("qualtrics_data") or s.get("linkedin_data") or s.get("clearinghouse_data")
    )]
    survey_count = len(with_survey)

    # ── Career outcomes — all sources ─────────────────────────────────────────
    outcomes = defaultdict(int)
    not_seeking_count = 0

    for s in students:
        outcome = _student_outcome(s)
        if outcome == "NOT seeking":
            not_seeking_count += 1
        else:
            outcomes[outcome] += 1

    # Knowledge rate: students with ANY resolved outcome (excl. Unresolved) / total graduates
    # Includes NOT seeking — we know their outcome; excludes only students with no resolvable data.
    known_count = total_graduates - outcomes.get("Unresolved", 0)

    outcome_order = [
        "Continuing education",
        "Employed full-time",
        "Employed part-time",
        "Volunteering or service program",
        "Serving in the U.S. Armed Forces",
        "Starting a business",
        "Unplaced",
        "Unresolved",
    ]
    outcomes_table = []
    grand_total = 0
    for label in outcome_order:
        n = outcomes.get(label, 0)
        outcomes_table.append({"label": label, "n": n})
        grand_total += n
    outcomes_table.append({"label": "Grand Total", "n": grand_total})
    if not_seeking_count:
        outcomes_table.append({"label": "NOT seeking", "n": not_seeking_count})

    employed_count = outcomes.get("Employed full-time", 0) + outcomes.get("Employed part-time", 0)

    # % in workforce: Employed FT/PT + Volunteering + Military + Business
    # Denominator: grand_total (all resolved outcomes excl. NOT seeking, per spec)
    in_workforce_count = (
        outcomes.get("Employed full-time", 0) +
        outcomes.get("Employed part-time", 0) +
        outcomes.get("Volunteering or service program", 0) +
        outcomes.get("Serving in the U.S. Armed Forces", 0) +
        outcomes.get("Starting a business", 0)
    )

    # ── Students whose outcome is "employed" — for Qualtrics-specific stats ───
    employed_qualtrics = [
        s for s in with_survey
        if "employed" in s["qualtrics_data"][0]["payload"].get("STATUS", "").lower()
        and "seeking" not in s["qualtrics_data"][0]["payload"].get("STATUS", "").lower()
    ]

    # ── Nature / modality / status of positions — Qualtrics + LinkedIn (all sources) ────
    nature_counts   = defaultdict(int)
    field_counts    = defaultdict(int)
    modality_counts = defaultdict(int)
    emp_status_counts = defaultdict(int)

    _OTHER_SKIP = {"other - none of the above", "other – none of the above",
                   "other-none of the above", "none of the above", "other"}

    def _skip_other(val: str) -> bool:
        return val.lower().strip() in _OTHER_SKIP

    for s in students:
        outcome = _student_outcome(s)
        if "employed" not in outcome.lower():
            continue
        # Qualtrics
        if s.get("qualtrics_data"):
            p = s["qualtrics_data"][0]["payload"]
            nat = p.get("EMP_NATURE", "").strip()
            if nat and not _skip_other(nat):
                nature_counts[nat] += 1
            fld = p.get("EMP_FIELD", "").strip()
            if fld and not _skip_other(fld):
                field_counts[fld] += 1
            mod = p.get("EMP_JOBSITE", "").strip()
            if mod and not _skip_other(mod):
                modality_counts[mod] += 1
            status = p.get("EMP_TYPE", "").strip()
            if status and not _skip_other(status):
                emp_status_counts[status] += 1
        # LinkedIn
        if s.get("linkedin_data"):
            li  = s["linkedin_data"][0]["payload"]
            mod = (li.get("modality_(hybrid_etc.if_known)") or li.get("employment_modality") or "").strip()
            if mod and not _skip_other(mod):
                modality_counts[mod] += 1

    # ── Salary (Qualtrics-only) ───────────────────────────────────────────────
    salaries = []
    bonus_values = []   # raw EMP_BONUS amounts for median calculation
    bonus_list   = []   # full list fallback if median can't be computed
    full_time_salary_respondents = 0
    for s in employed_qualtrics:
        p = s["qualtrics_data"][0]["payload"]
        emp_type = p.get("EMP_TYPE", "").lower()
        sal      = (p.get("EMP_SAL_1") or p.get("EMP_SAL") or p.get("EMP_SALARY") or "").strip()
        if "full-time" in emp_type or "full time" in emp_type or "employee" in emp_type:
            full_time_salary_respondents += 1
            if sal:
                mid = _parse_salary_midpoint(sal)
                if mid:
                    salaries.append(mid)
        bonus = p.get("EMP_BONUS", "").strip()
        if bonus and bonus.lower() not in ("", "no", "0", "none"):
            bonus_list.append(bonus)
            # Try to parse a numeric value for median calculation
            try:
                nums = re.findall(r"[\d,]+", bonus)
                if nums:
                    bonus_values.append(float(nums[0].replace(",", "")))
            except (ValueError, IndexError):
                pass

    bonus_count = len(bonus_list)
    bonus_sorted = sorted(bonus_values)
    bonus_median = int(_percentile(bonus_sorted, 50)) if len(bonus_sorted) >= 3 else None

    salaries_sorted = sorted(salaries)
    salary_stats = None
    if len(salaries_sorted) >= 5:
        salary_stats = {
            "n_reported":      len(salaries_sorted),
            "n_full_time":     full_time_salary_respondents,
            "bonus_count":     bonus_count,
            "bonus_median":    bonus_median,
            "bonus_list":      bonus_list if bonus_median is None else [],
            "p25": int(_percentile(salaries_sorted, 25)),
            "p50": int(_percentile(salaries_sorted, 50)),
            "p75": int(_percentile(salaries_sorted, 75)),
        }

    # ── Employment search methods (Qualtrics-only) ────────────────────────────
    emp_how_counts     = defaultdict(int)
    emp_how_respondents = 0
    for s in employed_qualtrics:
        p = s["qualtrics_data"][0]["payload"]
        found_any = False
        for i in range(1, 13):
            v = p.get(f"EMP_HOW_{i}", "").strip()
            if v and v in EMP_HOW_LABEL:
                emp_how_counts[EMP_HOW_LABEL[v]] += 1
                found_any = True
        if found_any:
            emp_how_respondents += 1
    emp_how_table = sorted(emp_how_counts.items(), key=lambda x: -x[1])

    # ── Geographic distribution — FT/PT employed only, EMP_STATE primary ───────
    # Per spec: only graduates who reported full or part time employment;
    # based on EMP_STATE column (fall back to extracting from EMP_CITY1_1).
    geo_counts      = defaultdict(int)
    geo_respondents = 0

    for s in students:
        outcome = _student_outcome(s)
        # Per spec: FT/PT employed, military, and business starters
        if outcome not in ("Employed full-time", "Employed part-time",
                           "Serving in the U.S. Armed Forces", "Starting a business"):
            continue
        # Qualtrics: prefer EMP_STATE directly, fall back to EMP_CITY1_1 parsing
        if s.get("qualtrics_data"):
            p = s["qualtrics_data"][0]["payload"]
            emp_state = p.get("EMP_STATE", "").strip()
            if emp_state:
                geo_counts[_resolve_state(emp_state)] += 1
                geo_respondents += 1
            else:
                city_str = p.get("EMP_CITY1_1", "").strip()
                if city_str:
                    geo_counts[_extract_state_from_city_str(city_str)] += 1
                    geo_respondents += 1
        # LinkedIn: employer_state + employer_country
        elif s.get("linkedin_data"):
            li = s["linkedin_data"][0]["payload"]
            state   = (li.get("employer_state")   or "").strip()
            country = (li.get("employer_country") or "").strip()
            if state:
                geo_counts[_resolve_state(state, country)] += 1
                geo_respondents += 1

    geo_table = sorted(geo_counts.items(), key=lambda x: -x[1])

    # ── Starting a business — Qualtrics + LinkedIn (all sources) ─────────────
    biz_details = []

    for s in students:
        # Qualtrics
        if s.get("qualtrics_data"):
            p = s["qualtrics_data"][0]["payload"]
            if "business" in p.get("STATUS", "").lower():
                org     = p.get("STBUS_ORG",     "").strip()
                purpose = p.get("STBUS_PURPOSE", "").strip()
                biz_details.append({"org": org or "N/A", "purpose": purpose or "N/A"})
        # LinkedIn
        if s.get("linkedin_data"):
            li = s["linkedin_data"][0]["payload"]
            biz_name = (li.get("name_of_started_business") or "").strip()
            biz_desc = (li.get("started_business_description") or "").strip()
            if biz_name:
                biz_details.append({"org": biz_name, "purpose": biz_desc or "N/A"})

    biz_count = outcomes.get("Starting a business", 0)

    # ── Volunteer / service — Qualtrics + LinkedIn (all sources) ─────────────
    vol_details = []

    for s in students:
        # Qualtrics
        if s.get("qualtrics_data"):
            p = s["qualtrics_data"][0]["payload"]
            if "service" in p.get("STATUS", "").lower() or "volunteer" in p.get("STATUS", "").lower():
                org  = (p.get("VOL_ORG", "") or p.get("VOL_ORG_1", "")).strip()
                role = p.get("VOL_ROLE",  "").strip()
                vol_details.append({"org": org or "N/A", "role": role or "N/A"})
        # LinkedIn
        if s.get("linkedin_data"):
            li = s["linkedin_data"][0]["payload"]
            vol_org  = (li.get("volunteer_organization") or "").strip()
            vol_role = (li.get("volunteer_role")         or "").strip()
            if vol_org:
                vol_details.append({"org": vol_org, "role": vol_role or "N/A"})

    vol_count = outcomes.get("Volunteering or service program", 0)
    mil_count = outcomes.get("Serving in the U.S. Armed Forces", 0)

    # ── Continuing education — Qualtrics + Clearinghouse + LinkedIn (all sources)
    cont_edu_count     = outcomes.get("Continuing education", 0)
    cont_edu_umd_count = 0
    degree_counts      = defaultdict(int)
    cont_edu_programs  = []

    _UNSPECIFIED = {"unspecified", "unknown", "n/a", "na", "none", "", "other"}

    def _blank_or_unspecified(val: str) -> bool:
        return not val or val.lower().strip() in _UNSPECIFIED

    def _add_ce(inst: str, prog: str, deg: str):
        nonlocal cont_edu_umd_count
        inst_clean = inst.split(",")[0].strip() if inst else ""
        if inst and "university of maryland" in inst.lower() and "college park" in inst.lower():
            cont_edu_umd_count += 1
        deg_val = deg if not _blank_or_unspecified(deg) else None
        if deg_val:
            degree_counts[deg_val] += 1
        inst_val = inst_clean if not _blank_or_unspecified(inst_clean) else None
        prog_val = prog if not _blank_or_unspecified(prog) else None
        if inst_val or prog_val:
            cont_edu_programs.append({
                "institution": inst_val or "",
                "program":     prog_val or "",
                "degree":      deg_val  or "",
            })

    for s in students:
        # Qualtrics CE entries
        if s.get("qualtrics_data"):
            p = s["qualtrics_data"][0]["payload"]
            if _map_status(p.get("STATUS", "").strip()) == "Continuing education":
                inst = p.get("CONTEDU_INST_1", "").strip()
                prog = p.get("CONTEDU_PROGRAM", "").strip()
                deg  = p.get("CONTEDU_DEGREE",  "").strip()
                if inst or prog or deg:
                    _add_ce(inst, prog, deg)
        # Clearinghouse CE entries
        if s.get("clearinghouse_data"):
            ch = s["clearinghouse_data"][0]["payload"]
            inst = (
                ch.get("College Name") or ch.get("Institution Name") or ch.get("institution") or ""
            ).strip()
            prog = (
                ch.get("Enrollment Major 1") or ch.get("enrollment_major_1") or
                ch.get("CIP Description") or ch.get("program") or ""
            ).strip()
            deg = (
                ch.get("Degree Title") or ch.get("degree_title") or
                ch.get("Credential Level") or ch.get("degree") or ""
            ).strip()
            if inst or prog:
                _add_ce(inst, prog, deg)
        # LinkedIn CE entries
        if s.get("linkedin_data"):
            if _student_outcome(s) == "Continuing education":
                li = s["linkedin_data"][0]["payload"]
                inst = (li.get("continuing_education_institution") or "").strip()
                prog = (li.get("continuing_education_program")     or "").strip()
                deg  = (li.get("continuing_education_degree")      or "").strip()
                if inst or prog:
                    _add_ce(inst, prog, deg)

    degree_table = sorted(degree_counts.items(), key=lambda x: -x[1])

    # ── Out-of-classroom experience (Qualtrics-only) ──────────────────────────
    otherexp_counts     = defaultdict(int)
    otherexp_respondents = 0
    for s in with_survey:
        p = s["qualtrics_data"][0]["payload"]
        found_any = False
        for field, label in OTHEREXP_LABEL.items():
            v = p.get(field, "")
            if v and str(v).strip() not in ("", "0"):
                otherexp_counts[label] += 1
                found_any = True
        if found_any:
            otherexp_respondents += 1
    otherexp_table = sorted(otherexp_counts.items(), key=lambda x: -x[1])

    # ── Internship participation (Qualtrics-only) ─────────────────────────────
    intern_respondents      = 0
    interns_with_any        = 0
    intern_two_plus         = 0
    intern_paid_count       = 0
    intern_credit_count     = 0
    intern_paid_students    = 0   # students with ≥1 paid internship
    intern_credit_students  = 0   # students with ≥1 credit internship
    total_internships_reported = 0
    intern_hourly_wages     = []
    intern_list             = []

    for s in with_survey:
        p = s["qualtrics_data"][0]["payload"]
        nin_str = p.get("NUMINTERN", "").strip()
        if nin_str == "":
            continue
        intern_respondents += 1
        try:
            nin = int(nin_str)
        except ValueError:
            continue

        if nin > 0:
            interns_with_any += 1
            total_internships_reported += nin
            if nin >= 2:
                intern_two_plus += 1

        student_had_paid   = False
        student_had_credit = False
        for i in range(1, nin + 1):
            org     = p.get(f"{i}_INT_ORG_1",  "").strip()
            title   = p.get(f"{i}_INT_TITLE",  "").strip()
            paid    = p.get(f"{i}_INT_PAID",   "").strip()
            credit  = p.get(f"{i}_INT_CREDIT", "").strip()
            howmuch = p.get(f"{i}_INT_HOWMUCH","").strip()

            is_paid   = paid.lower().startswith("yes") if paid else False
            is_credit = credit.lower() == "yes" if credit else False

            if is_paid:
                intern_paid_count += 1
                student_had_paid = True
            if is_credit:
                intern_credit_count += 1
                student_had_credit = True
            if howmuch:
                try:
                    intern_hourly_wages.append(float(howmuch))
                except ValueError:
                    pass
            if org or title:
                intern_list.append({
                    "org":    org.split(",")[0].strip() if org else "Unknown",
                    "title":  title or "Unknown",
                    "paid":   "Paid" if is_paid else "Unpaid",
                    "credit": "Yes" if is_credit else "No",
                })
        if student_had_paid:   intern_paid_students += 1
        if student_had_credit: intern_credit_students += 1

    intern_avg_wage = intern_med_wage = None
    if intern_hourly_wages:
        intern_avg_wage = sum(intern_hourly_wages) / len(intern_hourly_wages)
        intern_med_wage = _percentile(sorted(intern_hourly_wages), 50)

    # ── Appendix A: Employers — filter to FT/PT employed; EMP_ORG + EMP_TITLES ─
    _BLANK_SET = {"unspecified", "unknown", "n/a", "na", "none", ""}

    def _clean_field(val: str) -> str | None:
        v = val.strip() if val else ""
        return v if v.lower() not in _BLANK_SET else None

    employer_positions = []
    _seen_employer_uids: set = set()

    for s in students:
        uid     = s.get("uid", "")
        outcome = _student_outcome(s)
        # Deduplicate: one entry per student (uid) — skip if already recorded
        if uid and uid in _seen_employer_uids:
            continue
        # Qualtrics: filter by STATUS to FT or PT employed
        if s.get("qualtrics_data"):
            p      = s["qualtrics_data"][0]["payload"]
            status = p.get("STATUS", "").strip().lower()
            if "employed full" in status or "employed part" in status:
                org   = _clean_field(p.get("EMP_ORG",   "") or p.get("EMP_ORG_1",   ""))
                title = _clean_field(p.get("EMP_TITLES", "") or p.get("EMP_TITLE", ""))
                if org and title:
                    employer_positions.append({
                        "employer": org.split(",")[0].strip(),
                        "title":    title,
                    })
                    if uid:
                        _seen_employer_uids.add(uid)
        # LinkedIn: outcome must be employed
        elif s.get("linkedin_data") and "employed" in outcome.lower():
            li    = s["linkedin_data"][0]["payload"]
            org   = _clean_field(li.get("name_of_employer") or "")
            title = _clean_field(li.get("job_title") or "")
            if org and title:
                employer_positions.append({
                    "employer": org.split(",")[0].strip(),
                    "title":    title,
                })
                if uid:
                    _seen_employer_uids.add(uid)

    employer_positions_sorted = sorted(employer_positions, key=lambda x: x["employer"])

    # ── Appendix B: CE programs — Qualtrics + Clearinghouse + LinkedIn ────────
    # Deduplicate by (institution, program) key
    _seen_ce: set = set()
    cont_edu_programs_deduped = []
    for prog in cont_edu_programs:
        key = (prog["institution"].lower(), prog["program"].lower())
        if key not in _seen_ce:
            _seen_ce.add(key)
            cont_edu_programs_deduped.append(prog)

    cont_edu_programs_sorted = sorted(
        cont_edu_programs_deduped,
        key=lambda x: (x["institution"], x["program"]),
    )

    # ── Build and return ───────────────────────────────────────────────────────
    return {
        "meta": {
            "major":        (", ".join(major_filter) if isinstance(major_filter, list) else major_filter) or "All Majors",
            "school":       school_filter or "All Schools",
            "term":         (", ".join(term_filter) if isinstance(term_filter, list) else term_filter) or "All Terms",
            "generated_at": datetime.now().isoformat(),
        },
        "totals": {
            "total_graduates":       total_graduates,
            "survey_count":          survey_count,
            "linkedin_count":        len(with_linkedin),
            "clearinghouse_count":   len(with_clearinghouse),
            "known_count":           known_count,
            "survey_response_rate":  round(survey_count / total_graduates * 100, 1) if total_graduates else 0,
            "knowledge_rate":        round(known_count  / total_graduates * 100, 1) if total_graduates else 0,
        },
        "outcomes": {
            "table":          outcomes_table,
            "grand_total":    grand_total,
            # Placement: Placed / (Placed + Unplaced + Unresolved), excl. NOT seeking
            "placement_rate": round(
                (grand_total - outcomes.get("Unplaced", 0) - outcomes.get("Unresolved", 0)) /
                max(grand_total, 1) * 100, 1
            ) if grand_total else 0,
            "employed_count":     employed_count,
            "employed_pct":       round(employed_count / max(grand_total, 1) * 100, 1) if grand_total else 0,
            # % in workforce: Employed FT/PT + Volunteering + Military + Business
            # Denominator = people with data about them, excl. NOT seeking and Unresolved
            "in_workforce_count": in_workforce_count,
            "in_workforce_pct":   round(
                in_workforce_count /
                max(grand_total - outcomes.get("Unresolved", 0), 1) * 100, 1
            ) if grand_total else 0,
        },
        "nature": {
            "respondents":       len(employed_qualtrics),
            "nature_counts":     dict(nature_counts),
            "field_counts":      dict(field_counts),
            "modality_counts":   dict(modality_counts),
            "emp_status_counts": dict(emp_status_counts),
        },
        "salary":    salary_stats,
        "emp_search": {
            "respondents": emp_how_respondents,
            "table":       emp_how_table,
        },
        "geography": {
            "respondents": geo_respondents,
            "table":       geo_table,
        },
        "business":  {"count": biz_count, "details": biz_details},
        "volunteer": {"count": vol_count, "details": vol_details},
        "military":  {"count": mil_count},
        "continuing_education": {
            "count":       cont_edu_count,
            "umd_count":   cont_edu_umd_count,
            "degree_table": degree_table,
            "programs":    cont_edu_programs_sorted,
        },
        "otherexp": {
            "respondents": otherexp_respondents,
            "table":       otherexp_table,
        },
        "internships": {
            "respondents":      intern_respondents,
            "with_any":         interns_with_any,
            "two_plus":         intern_two_plus,
            "paid_count":       intern_paid_count,
            "credit_count":     intern_credit_count,
            "paid_students":    intern_paid_students,
            "credit_students":  intern_credit_students,
            "total_reported":   total_internships_reported,
            "avg_hourly_wage":  round(intern_avg_wage, 2) if intern_avg_wage else None,
            "median_hourly_wage": round(intern_med_wage, 2) if intern_med_wage else None,
            "intern_list":      intern_list,
        },
        "appendix_a": employer_positions_sorted,
        "appendix_b": cont_edu_programs_sorted,
    }


# ── DOCX helpers ──────────────────────────────────────────────────────────────

UMD_RED        = RGBColor(0xE2, 0x18, 0x33)
UMD_DARK       = RGBColor(0x1E, 0x29, 0x3B)
TABLE_HEADER_BG = "E21833"
TABLE_ROW_ALT   = "FFF3F3"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CCCCCC")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _header_row(table, headers, widths=None):
    row = table.rows[0]
    for i, (cell, text) in enumerate(zip(row.cells, headers)):
        cell.text = text
        _set_cell_bg(cell, TABLE_HEADER_BG)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        if widths:
            cell.width = Inches(widths[i])


def _data_row(table, row_idx, values, alt=True):
    row = table.rows[row_idx]
    bg  = TABLE_ROW_ALT if (alt and row_idx % 2 == 0) else "FFFFFF"
    for cell, val in zip(row.cells, values):
        cell.text = str(val)
        if bg != "FFFFFF":
            _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if row.cells[0].text == "Grand Total":
            for run in p.runs:
                run.bold = True


def _section_heading(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = UMD_RED
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after  = Pt(6)
    return h


def _body_text(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(text)
    p.style = "Normal"
    p.paragraph_format.space_after = Pt(6)
    if bold or italic:
        for run in p.runs:
            run.bold   = bold
            run.italic = italic
    return p


def _bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    return p


# ── DOCX generation ───────────────────────────────────────────────────────────

def generate_report_docx(data: dict) -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    meta        = data["meta"]
    totals      = data["totals"]
    outcomes    = data["outcomes"]
    nature      = data["nature"]
    salary      = data["salary"]
    emp_search  = data["emp_search"]
    geography   = data["geography"]
    biz         = data["business"]
    vol         = data["volunteer"]
    mil         = data["military"]
    cont_edu    = data["continuing_education"]
    otherexp    = data["otherexp"]
    internships = data["internships"]
    appendix_a  = data["appendix_a"]
    appendix_b  = data["appendix_b"]

    term_label  = meta["term"]
    if term_label and term_label != "All Terms" and len(term_label) >= 4 and term_label[:4].isdigit():
        year = int(term_label[:4])
    else:
        year = datetime.now().year
    major_label = meta["major"]

    # ── Cover ──────────────────────────────────────────────────────────────────
    title_p = doc.add_heading(f"Graduation Survey Report {year}", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_p.runs:
        run.font.color.rgb = UMD_RED

    sub = doc.add_paragraph(f"Report for {major_label} Graduates")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(4)
    for run in sub.runs:
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = UMD_DARK

    author_p = doc.add_paragraph("University Career Center & The President's Promise")
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in author_p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── Introduction ───────────────────────────────────────────────────────────
    _section_heading(doc, "Introduction")
    _body_text(doc,
        "The University Career Center tracks the initial destinations of UMD graduates through "
        "the Graduation Survey, which is administered to bachelor's degree recipients for each "
        "graduation cohort. The survey captures career-related outcomes, as well as data on "
        "participation in co-curricular activities during undergraduate studies. Data is also "
        "collected through LinkedIn profiles and the National Student Clearinghouse.")
    _body_text(doc,
        f"This report presents outcomes for {major_label} graduates"
        + (f" for term {term_label}" if term_label != "All Terms" else "") + ".")
    _body_text(doc,
        "Note: Throughout this report, percents may not sum to 100% due to rounding.",
        italic=True)

    # ── Response Rates ─────────────────────────────────────────────────────────
    _section_heading(doc, "Response Rates")
    _body_text(doc,
        f"As of {datetime.now().strftime('%B %Y')}, data from {totals['known_count']} of "
        f"{totals['total_graduates']} graduates have been collected via the survey or other means.")
    _bullet(doc, f"Survey response rate: {totals['survey_response_rate']}% "
                 f"({totals['survey_count']} of {totals['total_graduates']} graduates)")
    _bullet(doc, f"Knowledge rate: {totals['knowledge_rate']}% "
                 f"({totals['known_count']} of {totals['total_graduates']} graduates)")
    sources_used = []
    if totals["survey_count"]:
        sources_used.append(f"{totals['survey_count']} via Qualtrics survey")
    if totals.get("linkedin_count"):
        sources_used.append(f"{totals['linkedin_count']} via LinkedIn")
    if totals.get("clearinghouse_count"):
        sources_used.append(f"{totals['clearinghouse_count']} via National Student Clearinghouse")
    if sources_used:
        _body_text(doc, "Data sources: " + "; ".join(sources_used) + ".", italic=True)

    # ── Reported Outcomes ──────────────────────────────────────────────────────
    _section_heading(doc, "Reported Outcomes for Graduates")
    _bullet(doc, f"Placement rate: {outcomes['placement_rate']}%")
    grand = outcomes["grand_total"]
    unplaced_n    = next((r["n"] for r in outcomes["table"] if r["label"] == "Unplaced"),    0)
    unresolved_n  = next((r["n"] for r in outcomes["table"] if r["label"] == "Unresolved"), 0)
    if grand:
        _bullet(doc, f"Unplaced and unresolved: "
                     f"{_pct(unplaced_n, grand)} of students were unplaced, and "
                     f"{_pct(unresolved_n, grand)} were unresolved.")
        _bullet(doc, f"Entering the workforce: "
                     f"{_pct(outcomes['employed_count'], grand)} of graduates reported joining the workforce.")

    doc.add_paragraph()

    tbl = doc.add_table(rows=len(outcomes["table"]) + 1, cols=3)
    _add_table_borders(tbl)
    tbl.style = "Table Grid"
    _header_row(tbl, ["Career Outcomes", "N", "%"])
    for idx, row in enumerate(outcomes["table"]):
        n       = row["n"]
        pct_str = _pct(n, grand) if row["label"] != "Grand Total" else ""
        _data_row(tbl, idx + 1, [row["label"], str(n), pct_str])
        if row["label"] == "Grand Total":
            for cell in tbl.rows[idx + 1].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

    # ── Nature of Positions ────────────────────────────────────────────────────
    has_nature = (nature["nature_counts"] or nature["field_counts"]
                  or nature.get("modality_counts") or nature.get("emp_status_counts"))
    if nature["respondents"] >= 3 and has_nature:
        _section_heading(doc, "Nature of Positions")
        _body_text(doc,
            f"Out of the {nature['respondents']} graduates who answered survey questions about "
            "the nature of their employment:")
        if nature["nature_counts"]:
            total_nat = sum(nature["nature_counts"].values())
            for label, n in sorted(nature["nature_counts"].items(), key=lambda x: -x[1]):
                _bullet(doc, f"{_pct(n, total_nat)} – {label}")
        if nature["field_counts"]:
            total_fld = sum(nature["field_counts"].values())
            for label, n in sorted(nature["field_counts"].items(), key=lambda x: -x[1]):
                _bullet(doc, f"{_pct(n, total_fld)} – {label}")
        if nature.get("emp_status_counts"):
            total_status = sum(nature["emp_status_counts"].values())
            _body_text(doc, "Employment status:")
            for label, n in sorted(nature["emp_status_counts"].items(), key=lambda x: -x[1]):
                _bullet(doc, f"{_pct(n, total_status)} – {label}")
        if nature.get("modality_counts"):
            total_mod = sum(nature["modality_counts"].values())
            _body_text(doc, "Employment modality (in-person, remote, hybrid):")
            for label, n in sorted(nature["modality_counts"].items(), key=lambda x: -x[1]):
                _bullet(doc, f"{_pct(n, total_mod)} – {label}")

    # ── Salary ─────────────────────────────────────────────────────────────────
    if salary:
        _section_heading(doc, "Salary")
        _body_text(doc,
            "Graduates reported salaries and bonuses in ranges, so percentiles are calculated "
            "from grouped frequency data. Data is only provided if 5 or more graduates provided "
            "their salary ranges.")
        _body_text(doc,
            f"{salary['n_full_time']} graduates entering full-time employment reported their "
            f"starting salary ranges. {salary['bonus_count']} reported receiving a bonus.")
        sal_tbl = doc.add_table(rows=4, cols=2)
        _add_table_borders(sal_tbl)
        sal_tbl.style = "Table Grid"
        _header_row(sal_tbl, ["Percentile", "Salary"])
        _data_row(sal_tbl, 1, ["25th percentile",            f"${salary['p25']:,}"])
        _data_row(sal_tbl, 2, ["50th percentile (median)",   f"${salary['p50']:,}"])
        _data_row(sal_tbl, 3, ["75th percentile",            f"${salary['p75']:,}"])
        doc.add_paragraph()

    # ── Employment Search ──────────────────────────────────────────────────────
    if emp_search["respondents"] >= 3 and emp_search["table"]:
        _section_heading(doc, "Employment Search")
        _body_text(doc, "See Appendix A for a list of all self-reported employers and positions.")
        _body_text(doc,
            f"{emp_search['respondents']} respondents answered survey questions about how they "
            f"found their job. Their responses are below:")
        tbl2 = doc.add_table(rows=len(emp_search["table"]) + 1, cols=3)
        _add_table_borders(tbl2)
        tbl2.style = "Table Grid"
        _header_row(tbl2, ["Method Used to Find Employment", "N", "%"])
        for idx, (label, n) in enumerate(emp_search["table"]):
            _data_row(tbl2, idx + 1, [label, str(n), _pct(n, emp_search["respondents"])])
        _body_text(doc, "Note: Respondents could check all methods that applied.", italic=True)
        doc.add_paragraph()

    # ── Geographic Distribution ────────────────────────────────────────────────
    if geography["respondents"] >= 3 and geography["table"]:
        _section_heading(doc, "Geographic Distribution")
        total_geo = sum(n for _, n in geography["table"])
        md_n = next((n for s, n in geography["table"] if s == "Maryland"),              0)
        dc_n = next((n for s, n in geography["table"] if s == "District of Columbia"),  0)
        va_n = next((n for s, n in geography["table"] if s == "Virginia"),               0)
        _body_text(doc,
            f"{geography['respondents']} graduates who reported being employed listed where "
            f"they were working.")
        if md_n:
            dmv_n = md_n + dc_n + va_n
            _body_text(doc,
                f"Of those respondents, {_pct(md_n, geography['respondents'])} reported working "
                f"in Maryland, {_pct(dc_n, geography['respondents'])} in District of Columbia, "
                f"and {_pct(va_n, geography['respondents'])} in Virginia. In total, "
                f"{_pct(dmv_n, geography['respondents'])} of graduates reported working in "
                f"Maryland, Virginia, or Washington, D.C.")
        _body_text(doc, "The top states where graduates reported working were:")
        geo_tbl = doc.add_table(rows=len(geography["table"]) + 2, cols=3)
        _add_table_borders(geo_tbl)
        geo_tbl.style = "Table Grid"
        _header_row(geo_tbl, ["State", "N", "%"])
        for idx, (state, n) in enumerate(geography["table"]):
            _data_row(geo_tbl, idx + 1, [state, str(n), _pct(n, total_geo)])
        gt_row = geo_tbl.rows[-1]
        for cell, val in zip(gt_row.cells, ["Grand Total", str(total_geo), "100%"]):
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.bold = True
        doc.add_paragraph()

    # ── Starting a Business ────────────────────────────────────────────────────
    _section_heading(doc, "Starting a Business/Organization")
    if biz["count"] == 0:
        _body_text(doc, "No graduates reported starting their own business or organization.")
    else:
        s = "s" if biz["count"] > 1 else ""
        _body_text(doc, f"{biz['count']} graduate{s} reported starting their own business or organization.")

    # ── Service / Volunteer ────────────────────────────────────────────────────
    _section_heading(doc, "Service/Volunteer Programs")
    if vol["count"] == 0:
        _body_text(doc, "No graduates reported participating in a service or volunteer program after graduation.")
    else:
        s = "s" if vol["count"] > 1 else ""
        _body_text(doc,
            f"{vol['count']} graduate{s} reported participating in a service or volunteer "
            f"program after graduation.")

    # ── Continuing Education ───────────────────────────────────────────────────
    _section_heading(doc, "Continuing Education")
    if cont_edu["count"] == 0:
        _body_text(doc, "No graduates reported plans to continue their education.")
    else:
        _body_text(doc,
            f"{cont_edu['count']} respondents reported they would continue their education after "
            f"graduation. See Appendix B for a full list of the programs that graduates reported attending.")
        if cont_edu["umd_count"]:
            _body_text(doc,
                f"Out of these {cont_edu['count']} graduates, {cont_edu['umd_count']} said they "
                f"were attending the University of Maryland, College Park.")
        _body_text(doc,
            "Note: Some graduates may be employed and continuing their education simultaneously; "
            "they are not included here.", italic=True)

    if cont_edu["degree_table"]:
        doc.add_paragraph()
        deg_tbl   = doc.add_table(rows=len(cont_edu["degree_table"]) + 2, cols=3)
        _add_table_borders(deg_tbl)
        deg_tbl.style = "Table Grid"
        _header_row(deg_tbl, ["Degree Type", "N", "%"])
        total_deg = sum(n for _, n in cont_edu["degree_table"])
        for idx, (deg, n) in enumerate(cont_edu["degree_table"]):
            _data_row(deg_tbl, idx + 1, [deg, str(n), _pct(n, total_deg)])
        gt_row = deg_tbl.rows[-1]
        for cell, val in zip(gt_row.cells, ["Total", str(total_deg), "100.0%"]):
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.bold = True
        doc.add_paragraph()

    # ── Out of Classroom Experience ────────────────────────────────────────────
    _section_heading(doc, "Out of Classroom Experience")
    if otherexp["respondents"]:
        _body_text(doc,
            f"{otherexp['respondents']} respondents answered survey questions about their "
            f"experiences outside the classroom at UMD. Their responses are below:")
        if otherexp["table"]:
            exp_tbl = doc.add_table(rows=len(otherexp["table"]) + 1, cols=3)
            _add_table_borders(exp_tbl)
            exp_tbl.style = "Table Grid"
            _header_row(exp_tbl, ["Experience", "N", "%"])
            for idx, (label, n) in enumerate(otherexp["table"]):
                _data_row(exp_tbl, idx + 1, [label, str(n), _pct(n, otherexp["respondents"])])
            _body_text(doc, "Note: Respondents could check all responses that applied.", italic=True)
        doc.add_paragraph()
    else:
        _body_text(doc, "Insufficient data to report out-of-classroom experiences.")

    # ── Internship Participation ───────────────────────────────────────────────
    _section_heading(doc, "Internship Participation")
    if internships["respondents"]:
        _body_text(doc,
            f"{internships['respondents']} respondents answered questions about their internship "
            f"experiences while at UMD.")
        doc.add_paragraph()
        pct_any = _pct(internships["with_any"], internships["respondents"])
        _bullet(doc,
            f"{pct_any} of respondents ({internships['with_any']}) reported having at least one "
            f"internship during their time at the University of Maryland.")
        if internships["with_any"] and internships["two_plus"]:
            _bullet(doc,
                f"Among those graduates who reported having any internships, "
                f"{_pct(internships['two_plus'], internships['with_any'])} completed two or more internships.")
        if internships["with_any"]:
            _body_text(doc, f"Out of the {internships['with_any']} respondents who had at least one internship:")
            _bullet(doc,
                f"Paid internships: {_pct(internships['paid_students'], totals['total_graduates'])} "
                f"of graduates ({internships['paid_students']}) had at least one paid internship.")
            _bullet(doc,
                f"Internships for academic credit: "
                f"{_pct(internships['credit_students'], totals['total_graduates'])} of graduates "
                f"({internships['credit_students']}) had at least one internship for academic credit.")
    else:
        _body_text(doc, "Insufficient data to report internship participation.")

    # ── Internship Experiences ─────────────────────────────────────────────────
    if internships["intern_list"]:
        _section_heading(doc, "Internship Experiences")
        _body_text(doc,
            f"Among the {internships['with_any']} graduates who reported having at least one "
            f"internship, a total of {internships['total_reported']} internship experiences were reported.")
        if internships["avg_hourly_wage"]:
            _body_text(doc,
                f"Of internships that paid an hourly wage, the average reported income was "
                f"${internships['avg_hourly_wage']:.2f} per hour, and the median reported income "
                f"was ${internships['median_hourly_wage']:.2f} per hour.")
        doc.add_paragraph()
        int_tbl = doc.add_table(rows=len(internships["intern_list"]) + 1, cols=4)
        _add_table_borders(int_tbl)
        int_tbl.style = "Table Grid"
        _header_row(int_tbl, ["Organization", "Position", "Paid", "Credit"])
        for idx, entry in enumerate(internships["intern_list"]):
            _data_row(int_tbl, idx + 1, [entry["org"], entry["title"], entry["paid"], entry["credit"]])
        doc.add_paragraph()

    # ── Appendix A ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    _section_heading(doc, "Appendix A: Employers and Positions")
    _body_text(doc,
        f"The table below shows the self-reported employers and positions for {major_label} "
        f"graduates (Qualtrics survey and LinkedIn). "
        f"This list only includes positions if both the employer and position were reported.")
    if appendix_a:
        app_a_tbl = doc.add_table(rows=len(appendix_a) + 1, cols=2)
        _add_table_borders(app_a_tbl)
        app_a_tbl.style = "Table Grid"
        _header_row(app_a_tbl, ["Employer", "Job Title"])
        for idx, entry in enumerate(appendix_a):
            _data_row(app_a_tbl, idx + 1, [entry["employer"], entry["title"]])
    else:
        _body_text(doc, "No employer data available for this cohort.")
    doc.add_paragraph()

    # ── Appendix B ─────────────────────────────────────────────────────────────
    _section_heading(doc, "Appendix B: Continuing Education Programs")
    _body_text(doc,
        "The table below shows the programs that graduates reported attending "
        "(Qualtrics survey, National Student Clearinghouse, and LinkedIn). "
        "This list only includes programs if the university or program were known.")
    if appendix_b:
        app_b_tbl = doc.add_table(rows=len(appendix_b) + 1, cols=3)
        _add_table_borders(app_b_tbl)
        app_b_tbl.style = "Table Grid"
        _header_row(app_b_tbl, ["Institution", "Program", "Degree"])
        for idx, entry in enumerate(appendix_b):
            _data_row(app_b_tbl, idx + 1, [entry["institution"], entry["program"], entry["degree"]])
    else:
        _body_text(doc, "No continuing education data available for this cohort.")

    # ── Serialise ──────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Dashboard longitudinal aggregation ────────────────────────────────────────

def aggregate_dashboard_data(
    major_filter=None,
    school_filter=None,
    term_filter=None,
):
    """Aggregate comprehensive dashboard data including per-term longitudinal trends."""
    import database as _db

    # Overall summary (all selected terms combined)
    overall = aggregate_report_data(major_filter, school_filter, term_filter)

    # Per-term longitudinal breakdowns
    all_terms = sorted(_db.get_distinct_terms())
    if term_filter:
        all_terms = [t for t in all_terms if t in term_filter]

    longitudinal = []
    for term in all_terms:
        td = aggregate_report_data(major_filter, school_filter, [term])
        total_grads = td["totals"]["total_graduates"]
        if total_grads == 0:
            continue
        sal  = td.get("salary") or {}
        intn = td.get("internships") or {}
        longitudinal.append({
            "term":                 term,
            "total_graduates":      total_grads,
            "placement_rate":       td["outcomes"]["placement_rate"],
            "knowledge_rate":       td["totals"]["knowledge_rate"],
            "survey_response_rate": td["totals"]["survey_response_rate"],
            "in_workforce_pct":     td["outcomes"]["in_workforce_pct"],
            "salary_p25":           sal.get("p25"),
            "salary_p50":           sal.get("p50"),
            "salary_p75":           sal.get("p75"),
            "internship_pct":       round(intn.get("paid_students", 0) / total_grads * 100, 1),
            "outcomes":             {row["label"]: row["n"] for row in td["outcomes"]["table"]},
        })

    # Per-school comparison (skip when a school is already selected)
    school_comparison = []
    if not school_filter:
        for school in sorted(_db.get_distinct_values("major1_coll")):
            sd = aggregate_report_data(major_filter, school, term_filter)
            if sd["totals"]["total_graduates"] == 0:
                continue
            school_comparison.append({
                "school":           school,
                "total":            sd["totals"]["total_graduates"],
                "placement_rate":   sd["outcomes"]["placement_rate"],
                "knowledge_rate":   sd["totals"]["knowledge_rate"],
                "in_workforce_pct": sd["outcomes"]["in_workforce_pct"],
            })

    return {
        "overall":           overall,
        "longitudinal":      longitudinal,
        "school_comparison": school_comparison,
    }


def aggregate_major_comparison(
    major_filter=None,
    school_filter=None,
    term_filter=None,
):
    """Per-major outcome stats for the Major Analytics dashboard tab."""
    students = database.get_students_with_data(
        limit=None, offset=None,
        major_filter=major_filter,
        school_filter=school_filter,
        term_filter=term_filter,
    )

    from collections import defaultdict

    major_groups: dict = defaultdict(list)
    for s in students:
        major  = (s.get("major")  or "Unknown").strip()
        school = (s.get("school") or "").strip()
        major_groups[(major, school)].append(s)

    results = []
    for (major, school), group in sorted(major_groups.items(), key=lambda x: -len(x[1])):
        if len(group) < 3:
            continue

        outcomes_list   = [_student_outcome(s) for s in group]
        not_seeking_cnt = sum(1 for o in outcomes_list if o == "NOT seeking")
        known_cnt       = sum(1 for o in outcomes_list if o != "Unresolved")
        placed_cnt      = sum(1 for o in outcomes_list if o not in ("Unresolved", "Unplaced", "NOT seeking"))
        workforce_cnt   = sum(1 for o in outcomes_list if o in (
            "Employed full-time", "Employed part-time",
            "Serving in the U.S. Armed Forces", "Starting a business",
            "Volunteering or service program",
        ))

        salaries = []
        for s in group:
            if s.get("qualtrics_data"):
                p   = s["qualtrics_data"][0]["payload"]
                sal = (p.get("EMP_SAL_1") or p.get("EMP_SAL") or p.get("EMP_SALARY") or "").strip()
                if sal:
                    mid = _parse_salary_midpoint(sal)
                    if mid:
                        salaries.append(mid)

        sal_sorted = sorted(salaries)
        denom = max(known_cnt - not_seeking_cnt, 1)

        results.append({
            "major":           major,
            "school":          school,
            "total":           len(group),
            "placed":          placed_cnt,
            "known":           known_cnt,
            "in_workforce":    workforce_cnt,
            "placement_rate":  round(placed_cnt / denom * 100, 1),
            "knowledge_rate":  round(known_cnt  / len(group) * 100, 1),
            "in_workforce_pct": round(workforce_cnt / max(known_cnt, 1) * 100, 1),
            "salary_p25":      int(_percentile(sal_sorted, 25)) if len(sal_sorted) >= 3 else None,
            "salary_median":   int(_percentile(sal_sorted, 50)) if len(sal_sorted) >= 3 else None,
            "salary_p75":      int(_percentile(sal_sorted, 75)) if len(sal_sorted) >= 3 else None,
        })

    return sorted(results, key=lambda x: -x["total"])
